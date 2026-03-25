"""
OpenCaselist MCP Server v9

Changes over v8:
  - Rich formatting extraction: underline, bold, highlight preserved from source .docx
  - Output cards as a downloadable .docx via /cards/{doc_id} endpoint
  - Formatting follows NDT/CEDA verbatim spec (cyan highlight + underline)
  - Plain-text fallback still included in tool response for readability

Flow:
  1. GET /v1/search?q=...&shard=...        → matched paths (team + file)
  2. Download each matched .docx via /v1/download?path=<path>  (authed, cached)
  3. Parse .docx preserving run-level formatting (bold/underline/highlight)
  4. Find card blocks matching query terms
  5. Deduplicate repeated cards across files
  6. Build output .docx with verbatim formatting, store server-side
  7. Return tool response text + download URL for the .docx

Required Render environment variables:
  TABROOM_USERNAME   — your Tabroom.com email
  TABROOM_PASSWORD   — your Tabroom.com password
  CASELIST_SHARD     — e.g. "ndtceda25" (default)
  SERVER_BASE_URL    — e.g. "https://your-app.onrender.com" (for docx download links)
"""

import json
import asyncio
import hashlib
import os
import uuid
import io
import httpx
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import AsyncGenerator

from fastapi import FastAPI, Request
from fastapi.responses import StreamingResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

# ─────────────────────────────────────────────
# Config / tuning knobs
# ─────────────────────────────────────────────

TABROOM_USERNAME = os.environ.get("TABROOM_USERNAME", "")
TABROOM_PASSWORD = os.environ.get("TABROOM_PASSWORD", "")
CASELIST_SHARD   = os.environ.get("CASELIST_SHARD", "ndtceda25")
SERVER_BASE_URL  = os.environ.get("SERVER_BASE_URL", "https://jioajfoeaijfiaeojf-3.onrender.com")
API_BASE         = "https://api.opencaselist.com/v1"

MAX_TEAMS        = int(os.environ.get("MAX_TEAMS", "20"))
MAX_FILES        = int(os.environ.get("MAX_FILES", "40"))
FILE_TIMEOUT     = float(os.environ.get("FILE_TIMEOUT", "7"))
CONCURRENCY      = int(os.environ.get("CONCURRENCY", "8"))

# In-process download cache: path → bytes
_download_cache: dict[str, bytes] = {}

# In-process docx output store: doc_id → bytes (generated output docs)
_docx_store: dict[str, bytes] = {}

# Persistent authenticated session
_session: httpx.AsyncClient | None = None
_dl_semaphore: asyncio.Semaphore | None = None


def get_semaphore() -> asyncio.Semaphore:
    global _dl_semaphore
    if _dl_semaphore is None:
        _dl_semaphore = asyncio.Semaphore(CONCURRENCY)
    return _dl_semaphore


# ─────────────────────────────────────────────
# Auth / session
# ─────────────────────────────────────────────

async def get_session() -> httpx.AsyncClient | None:
    global _session
    if _session is None:
        _session = await _create_session()
    return _session


async def _create_session() -> httpx.AsyncClient | None:
    if not TABROOM_USERNAME or not TABROOM_PASSWORD:
        return None
    client = httpx.AsyncClient(follow_redirects=True, timeout=20.0)
    resp = await client.post(
        f"{API_BASE}/login",
        json={"username": TABROOM_USERNAME, "password": TABROOM_PASSWORD, "remember": True},
    )
    if resp.status_code in (200, 201) and "caselist_token" in client.cookies:
        return client
    await client.aclose()
    return None


async def authed_get(url: str, params: dict = None) -> httpx.Response:
    """GET with the persistent session; re-authenticates once on 401."""
    global _session
    session = await get_session()
    if session is None:
        raise RuntimeError("Cannot authenticate — check TABROOM_USERNAME/PASSWORD")
    resp = await session.get(url, params=params)
    if resp.status_code == 401:
        await session.aclose()
        _session = await _create_session()
        if _session is None:
            raise RuntimeError("Re-authentication failed")
        resp = await _session.get(url, params=params)
    return resp


# ─────────────────────────────────────────────
# Download (cached, timeout-guarded, semaphored)
# ─────────────────────────────────────────────

async def download_doc(path: str) -> bytes | None:
    if not path:
        return None

    if path in _download_cache:
        return _download_cache[path]

    async with get_semaphore():
        if path in _download_cache:
            return _download_cache[path]

        try:
            session = await get_session()
            if session is None:
                return None

            resp = await asyncio.wait_for(
                session.get(f"{API_BASE}/download", params={"path": path}),
                timeout=FILE_TIMEOUT,
            )
            if resp.status_code == 401:
                global _session
                await session.aclose()
                _session = await _create_session()
                session = _session
                if session is None:
                    return None
                resp = await asyncio.wait_for(
                    session.get(f"{API_BASE}/download", params={"path": path}),
                    timeout=FILE_TIMEOUT,
                )

            if resp.status_code == 200 and len(resp.content) > 500 and resp.content[:2] == b'PK':
                _download_cache[path] = resp.content
                return resp.content

            if path.startswith("http"):
                resp2 = await asyncio.wait_for(
                    session.get(path),
                    timeout=FILE_TIMEOUT,
                )
                if resp2.status_code == 200 and len(resp2.content) > 500 and resp2.content[:2] == b'PK':
                    _download_cache[path] = resp2.content
                    return resp2.content

        except (asyncio.TimeoutError, Exception):
            pass

    return None


# ─────────────────────────────────────────────
# Rich Docx parsing — preserves run-level formatting
# ─────────────────────────────────────────────

def _get_highlight_color(run) -> WD_COLOR_INDEX | None:
    """Safely extract highlight color from a run."""
    try:
        return run.font.highlight_color
    except Exception:
        return None


def _get_font_color(run) -> RGBColor | None:
    """Safely extract font color from a run."""
    try:
        clr = run.font.color
        if clr and clr.type is not None:
            return clr.rgb
    except Exception:
        pass
    return None


def extract_paragraphs_rich(content: bytes) -> list[dict]:
    """
    Extract paragraphs with full run-level formatting preserved.
    Returns a list of paragraph dicts:
      {
        "text": str,            — plain text of the full paragraph
        "style": str,           — paragraph style name (e.g. "Heading 4", "Normal")
        "runs": [
          {
            "text": str,
            "bold": bool,
            "italic": bool,
            "underline": bool,
            "highlight_color": WD_COLOR_INDEX | None,
            "font_color": RGBColor | None,
          }, ...
        ]
      }
    """
    try:
        doc = docx.Document(io.BytesIO(content))
        paragraphs = []
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            style_name = "Normal"
            try:
                style_name = p.style.name
            except Exception:
                pass

            runs = []
            for run in p.runs:
                if not run.text:
                    continue
                runs.append({
                    "text": run.text,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "underline": bool(run.underline),
                    "highlight_color": _get_highlight_color(run),
                    "font_color": _get_font_color(run),
                })
            if runs:
                paragraphs.append({
                    "text": p.text.strip(),
                    "style": style_name,
                    "runs": runs,
                })
        return paragraphs
    except Exception as e:
        return [{"text": f"[docx parse error: {e}]", "style": "Normal", "runs": [{"text": f"[docx parse error: {e}]", "bold": False, "italic": False, "underline": False, "highlight_color": None, "font_color": None}]}]


def find_card_blocks_rich(
    paragraphs: list[dict],
    query_terms: list[str],
    context: int = 20,
) -> list[list[dict]]:
    """
    Return windows of rich paragraph dicts around query term hits.
    Returns a list of blocks; each block is a list of paragraph dicts.
    """
    hits: set[int] = set()
    for i, para in enumerate(paragraphs):
        if any(term in para["text"].lower() for term in query_terms):
            for j in range(max(0, i - 2), min(len(paragraphs), i + context)):
                hits.add(j)

    if not hits:
        return []

    sorted_hits = sorted(hits)
    blocks, start, prev = [], sorted_hits[0], sorted_hits[0]
    for idx in sorted_hits[1:]:
        if idx > prev + 1:
            blocks.append((start, prev))
            start = idx
        prev = idx
    blocks.append((start, prev))
    return [paragraphs[s:e + 1] for s, e in blocks]


def dedup_cards(cards: list[dict]) -> list[dict]:
    seen, result = set(), []
    for card in cards:
        fp = hashlib.md5(card["text"][:300].strip().encode()).hexdigest()
        if fp not in seen:
            seen.add(fp)
            result.append(card)
    return result


# ─────────────────────────────────────────────
# Verbatim .docx output builder
# ─────────────────────────────────────────────

# Verbatim cyan: matches the standard NDT/CEDA turquoise highlight
# WD_COLOR_INDEX uses TURQUOISE (not CYAN) for the teal/cyan color
VERBATIM_HIGHLIGHT = WD_COLOR_INDEX.TURQUOISE


def _ensure_style(doc_obj, style_name: str, base_name: str = "Normal"):
    """Return a style by name, creating it if it doesn't exist."""
    try:
        return doc_obj.styles[style_name]
    except KeyError:
        style = doc_obj.styles.add_style(style_name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc_obj.styles[base_name]
        return style


def build_cards_docx(all_cards: list[dict]) -> bytes:
    """
    Build a verbatim-formatted .docx from a list of card dicts.

    Each card must have:
      "team"            — e.g. "ndtceda25/Harvard/AB"
      "round"           — e.g. "AB (Harvard) — Tournament Rd 2 Aff"
      "source_url"      — file path or URL (may be empty)
      "title"           — card block title / tag
      "text"            — plain text (used for fallback)
      "rich_paragraphs" — list of paragraph dicts from extract_paragraphs_rich
    """
    out = docx.Document()

    # Page setup — US Letter, narrow margins
    sec = out.sections[0]
    sec.page_width  = int(8.5 * 914400)  # EMU
    sec.page_height = int(11  * 914400)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)

    # Document-level font
    out.styles["Normal"].font.name = "Arial"
    out.styles["Normal"].font.size = Pt(10)

    for i, card in enumerate(all_cards):
        # ── Team / source header ──────────────────────────────────────────
        header_para = out.add_paragraph()
        try:
            header_para.style = out.styles["Heading 3"]
        except KeyError:
            pass
        run = header_para.add_run(f"[{i+1}] {card['team']}  |  {card['round']}")
        run.bold = True

        if card.get("source_url"):
            url_para = out.add_paragraph()
            url_run = url_para.add_run(f"Source: {card['source_url']}")
            url_run.font.size = Pt(8)
            try:
                url_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            except Exception:
                pass

        # ── Card title / tag ──────────────────────────────────────────────
        tag_para = out.add_paragraph()
        try:
            tag_para.style = out.styles["Heading 4"]
        except KeyError:
            pass
        tag_run = tag_para.add_run(card.get("title", "Untitled"))
        tag_run.bold = True

        # ── Rich paragraphs (preserving source formatting) ────────────────
        rich_paras = card.get("rich_paragraphs", [])

        if rich_paras:
            for para_data in rich_paras:
                p = out.add_paragraph()

                # Apply paragraph style if it maps to a known docx style
                src_style = para_data.get("style", "Normal")
                try:
                    p.style = out.styles[src_style]
                except KeyError:
                    pass  # keep Normal

                for run_data in para_data["runs"]:
                    r = p.add_run(run_data["text"])
                    r.bold      = run_data.get("bold", False)
                    r.italic    = run_data.get("italic", False)
                    r.underline = run_data.get("underline", False)

                    # Preserve source highlight; default verbatim cards to cyan
                    src_highlight = run_data.get("highlight_color")
                    if src_highlight is not None:
                        try:
                            r.font.highlight_color = src_highlight
                        except Exception:
                            r.font.highlight_color = VERBATIM_HIGHLIGHT
                    elif run_data.get("underline"):
                        # Underlined text in verbatim gets cyan highlight per spec
                        r.font.highlight_color = VERBATIM_HIGHLIGHT

                    # Preserve font color if set
                    src_color = run_data.get("font_color")
                    if src_color is not None:
                        try:
                            r.font.color.rgb = src_color
                        except Exception:
                            pass
        else:
            # Fallback: plain text with verbatim formatting applied to whole block
            p = out.add_paragraph()
            r = p.add_run(card.get("text", ""))
            r.underline = True
            r.font.highlight_color = VERBATIM_HIGHLIGHT

        # ── Divider between cards ─────────────────────────────────────────
        div = out.add_paragraph()
        div_run = div.add_run("─" * 80)
        div_run.font.size = Pt(7)
        try:
            div_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        except Exception:
            pass

    buf = io.BytesIO()
    out.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# Parse helpers
# ─────────────────────────────────────────────

def parse_team_from_path(path: str) -> tuple[str, str, str] | None:
    parts = path.strip("/").split("/")
    if len(parts) >= 3:
        return parts[0], parts[1], parts[2]
    return None


def get_file_path_from_search_hit(path: str) -> str | None:
    if "." in path.split("/")[-1]:
        return path
    return None


# ─────────────────────────────────────────────
# Core search logic
# ─────────────────────────────────────────────

async def process_file(
    file_path: str,
    team_label: str,
    caselist: str, school: str, team: str,
    query_terms: list[str],
) -> list[dict]:
    """Download one .docx and return matching card blocks with rich formatting."""
    content = await download_doc(file_path)
    if content is None:
        return []

    rich_paragraphs = extract_paragraphs_rich(content)
    rich_blocks = find_card_blocks_rich(rich_paragraphs, query_terms)

    return [
        {
            "team": f"{caselist}/{school}/{team}",
            "round": team_label,
            "source_url": file_path,
            "title": f"Block {i + 1} — {team_label}",
            # plain text for dedup / readable response
            "text": "\n".join(p["text"] for p in block),
            # rich data for docx output
            "rich_paragraphs": block,
        }
        for i, block in enumerate(rich_blocks)
    ]


async def fetch_rounds_and_search(
    caselist: str, school: str, team: str,
    query_terms: list[str],
    known_file_path: str | None,
    file_budget: asyncio.Semaphore,
) -> list[dict]:
    team_label = f"{team} ({school})"

    if known_file_path:
        async with file_budget:
            return await process_file(known_file_path, team_label, caselist, school, team, query_terms)

    try:
        resp = await asyncio.wait_for(
            authed_get(f"{API_BASE}/caselists/{caselist}/schools/{school}/teams/{team}/rounds"),
            timeout=10.0,
        )
    except Exception:
        return []

    if resp.status_code != 200:
        return []

    rounds = resp.json()
    if not isinstance(rounds, list):
        return []

    rounds_with_docs = [r for r in rounds if r.get("opensource")]
    if not rounds_with_docs:
        return []

    async def do_round(rnd: dict) -> list[dict]:
        opensource_path = rnd.get("opensource", "")
        tournament = rnd.get("tournament", "?")
        side = rnd.get("side", "")
        round_num = rnd.get("round", "?")
        label = f"{team} ({school}) — {tournament} Rd {round_num} {side}".strip()

        try:
            await asyncio.wait_for(file_budget.acquire(), timeout=1.0)
        except asyncio.TimeoutError:
            return []

        try:
            return await process_file(opensource_path, label, caselist, school, team, query_terms)
        finally:
            file_budget.release()

    results = await asyncio.gather(*[do_round(r) for r in rounds_with_docs], return_exceptions=True)
    cards = []
    for r in results:
        if isinstance(r, list):
            cards.extend(r)
    return cards


async def fetch_cites_for_team(
    caselist: str, school: str, team: str,
    query_terms: list[str],
) -> list[dict]:
    try:
        resp = await asyncio.wait_for(
            authed_get(f"{API_BASE}/caselists/{caselist}/schools/{school}/teams/{team}/cites"),
            timeout=10.0,
        )
    except Exception:
        return []

    if resp.status_code != 200:
        return []

    cites = resp.json()
    if not isinstance(cites, list):
        return []

    return [
        {
            "team": f"{caselist}/{school}/{team}",
            "round": "cites",
            "source_url": "",
            "title": cite.get("title", "Untitled"),
            "text": cite.get("cites", ""),
            "rich_paragraphs": [],  # no source doc available for cites
        }
        for cite in cites
        if any(term in (cite.get("title", "") + cite.get("cites", "")).lower()
               for term in query_terms)
    ]


# ─────────────────────────────────────────────
# Main search orchestrator
# ─────────────────────────────────────────────

async def search_opencaselist(query: str, shard: str = None) -> list[dict]:
    target_shard = shard or CASELIST_SHARD
    query_terms  = [t.lower() for t in query.split() if len(t) > 1]

    session = await get_session()
    if session is None:
        return [{"team": "error", "round": "", "source_url": "",
                 "title": "Authentication Error",
                 "text": "Set TABROOM_USERNAME and TABROOM_PASSWORD in Render environment variables.",
                 "rich_paragraphs": []}]

    # ── Step 1: Search index ──────────────────────────────────────────────────
    try:
        resp = await asyncio.wait_for(
            authed_get(f"{API_BASE}/search", params={"q": query, "shard": target_shard}),
            timeout=15.0,
        )
    except asyncio.TimeoutError:
        return [{"team": "error", "round": "", "source_url": "",
                 "title": "Search API Timeout", "text": "The search index did not respond in time.",
                 "rich_paragraphs": []}]

    if resp.status_code != 200:
        return [{"team": "error", "round": "", "source_url": "",
                 "title": f"Search API Error {resp.status_code}", "text": resp.text[:300],
                 "rich_paragraphs": []}]

    data  = resp.json()
    items = data if isinstance(data, list) else data.get("results", [data])

    # ── Step 2: Collect unique teams + their known file paths ────────────────
    seen_teams: dict[str, str | None] = {}

    for item in items:
        path = item.get("path", "")
        parsed = parse_team_from_path(path)
        if not parsed:
            continue
        caselist, school, team = parsed
        key = f"{caselist}/{school}/{team}"
        file_path = get_file_path_from_search_hit(path)
        if key not in seen_teams:
            seen_teams[key] = file_path
        elif file_path and seen_teams[key] is None:
            seen_teams[key] = file_path

        if len(seen_teams) >= MAX_TEAMS:
            break

    if not seen_teams:
        return [{"team": "-", "round": "", "source_url": "",
                 "title": "No Results",
                 "text": f"No teams found for '{query}' in '{target_shard}'.",
                 "rich_paragraphs": []}]

    # ── Step 3: Download + search (concurrently, budget-capped) ──────────────
    file_budget = asyncio.Semaphore(MAX_FILES)

    async def handle_team(key: str, known_file: str | None) -> list[dict]:
        parts = key.split("/", 2)
        if len(parts) != 3:
            return []
        caselist, school, team = parts

        cards = await fetch_rounds_and_search(
            caselist, school, team, query_terms, known_file, file_budget
        )

        has_real = any(not c["title"].startswith(("⚠️", "📄", "❌")) for c in cards)
        if not has_real:
            cite_cards = await fetch_cites_for_team(caselist, school, team, query_terms)
            if cite_cards:
                return cite_cards

        return cards

    nested = await asyncio.gather(
        *[handle_team(key, fp) for key, fp in seen_teams.items()],
        return_exceptions=True,
    )

    all_cards: list[dict] = []
    for r in nested:
        if isinstance(r, list):
            all_cards.extend(r)

    deduped = dedup_cards(all_cards)

    if not deduped:
        team_names = " | ".join(seen_teams.keys())
        return [{"team": team_names, "round": "", "source_url": "",
                 "title": "Teams found but no cards extracted",
                 "text": f"Searched {len(seen_teams)} team(s): {team_names}",
                 "rich_paragraphs": []}]

    return deduped


# ─────────────────────────────────────────────
# SSE / MCP boilerplate
# ─────────────────────────────────────────────

_client_queues: dict[str, asyncio.Queue] = {}


async def sse_stream(client_id: str) -> AsyncGenerator[str, None]:
    queue: asyncio.Queue = asyncio.Queue()
    _client_queues[client_id] = queue
    yield f"event: endpoint\ndata: /messages?client_id={client_id}\n\n"
    try:
        while True:
            try:
                message = await asyncio.wait_for(queue.get(), timeout=30.0)
                yield f"data: {json.dumps(message)}\n\n"
            except asyncio.TimeoutError:
                yield ": ping\n\n"
    finally:
        _client_queues.pop(client_id, None)


def make_response(id, result):
    return {"jsonrpc": "2.0", "id": id, "result": result}

def make_error(id, code, message):
    return {"jsonrpc": "2.0", "id": id, "error": {"code": code, "message": message}}


async def handle_jsonrpc(message: dict) -> dict | None:
    method = message.get("method")
    msg_id = message.get("id")
    params = message.get("params", {})

    if method == "initialize":
        return make_response(msg_id, {
            "protocolVersion": "2024-11-05",
            "capabilities": {"tools": {}},
            "serverInfo": {"name": "opencaselist-mcp", "version": "9.0.0"},
        })

    if method == "notifications/initialized":
        return None

    if method == "tools/list":
        return make_response(msg_id, {
            "tools": [{
                "name": "search_caselist",
                "description": (
                    f"Search OpenCaselist for full debate evidence cards. Default shard: {CASELIST_SHARD}. "
                    "Downloads open-source round docs per team via /v1/download (authenticated), "
                    "searches inside each one preserving verbatim formatting (underline/highlight/bold), "
                    "deduplicates repeated cards, and falls back to /cites when no open-source docs exist. "
                    "Returns card text AND a download URL for a formatted .docx Word file with "
                    "NDT/CEDA verbatim formatting (cyan highlight + underline). "
                    f"Caps at {MAX_TEAMS} teams and {MAX_FILES} file downloads per query."
                ),
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Search terms, e.g. 'faciality'"},
                        "shard": {
                            "type": "string",
                            "description": f"Shard (default: {CASELIST_SHARD}). E.g. ndtceda24, ndtceda25, hspolicy25",
                        },
                    },
                    "required": ["query"],
                },
            }]
        })

    if method == "tools/call":
        tool_name = params.get("name")
        arguments = params.get("arguments", {})

        if tool_name == "search_caselist":
            query          = arguments.get("query", "")
            shard_override = arguments.get("shard")
            cards = await search_opencaselist(query, shard=shard_override)

            # ── Build output .docx ────────────────────────────────────────
            doc_id = str(uuid.uuid4())
            try:
                docx_bytes = build_cards_docx(cards)
                _docx_store[doc_id] = docx_bytes
                download_url = f"{SERVER_BASE_URL}/cards/{doc_id}.docx"
                docx_note = f"\n\n📥 **Download cards as Word doc:** {download_url}\n"
            except Exception as e:
                docx_note = f"\n\n⚠️ Could not build .docx: {e}\n"

            # ── Plain-text response (still readable in Claude) ────────────
            effective_shard = shard_override or CASELIST_SHARD
            lines = [f"**{len(cards)} result(s) for '{query}'** (shard: {effective_shard}):{docx_note}"]
            for card in cards:
                source = f"\n🔗 {card['source_url']}" if card.get("source_url") else ""
                lines.append(
                    f"---\n📁 **{card['team']}** | {card['round']}{source}\n"
                    f"🏷 {card['title']}\n\n{card['text']}\n"
                )

            return make_response(msg_id, {"content": [{"type": "text", "text": "\n".join(lines)}]})

        return make_error(msg_id, -32601, f"Unknown tool: {tool_name}")

    return make_error(msg_id, -32601, f"Method not found: {method}")


# ─────────────────────────────────────────────
# FastAPI app
# ─────────────────────────────────────────────

app = FastAPI(title="OpenCaselist MCP Server")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


@app.get("/sse")
async def sse_endpoint(request: Request):
    client_id = str(uuid.uuid4())
    return StreamingResponse(
        sse_stream(client_id),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"},
    )


@app.post("/messages")
async def messages_endpoint(request: Request):
    client_id = request.query_params.get("client_id")
    body = await request.json()
    messages = body if isinstance(body, list) else [body]
    for message in messages:
        response = await handle_jsonrpc(message)
        if response is None:
            continue
        if client_id and client_id in _client_queues:
            await _client_queues[client_id].put(response)
        else:
            return JSONResponse(content=response)
    return JSONResponse(content={"status": "ok"})


@app.get("/cards/{doc_id}.docx")
async def download_cards_docx(doc_id: str):
    """
    Serve a previously generated cards .docx by its UUID.
    Files are stored in-process memory — they persist for the server's lifetime.
    """
    docx_bytes = _docx_store.get(doc_id)
    if docx_bytes is None:
        return JSONResponse(status_code=404, content={"error": "Document not found or expired."})
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="caselist_cards_{doc_id[:8]}.docx"'},
    )


@app.get("/.well-known/oauth-protected-resource")
async def oauth_resource():
    return JSONResponse(content={
        "resource": SERVER_BASE_URL,
        "authorization_servers": [],
    })

@app.get("/.well-known/oauth-authorization-server")
async def oauth_auth_server():
    return JSONResponse(content={
        "issuer": SERVER_BASE_URL,
        "authorization_endpoint": "",
        "token_endpoint": "",
        "response_types_supported": [],
    })

@app.post("/register")
async def register():
    return JSONResponse(status_code=400, content={"error": "registration_not_supported"})

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "version": "9.0.0",
        "shard": CASELIST_SHARD,
        "credentials_set": bool(TABROOM_USERNAME and TABROOM_PASSWORD),
        "docx_store_size": len(_docx_store),
        "limits": {"max_teams": MAX_TEAMS, "max_files": MAX_FILES,
                   "file_timeout_s": FILE_TIMEOUT, "concurrency": CONCURRENCY},
    }

@app.get("/")
async def root():
    return {"name": "OpenCaselist MCP Server", "version": "9.0.0", "shard": CASELIST_SHARD}


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    print(f"🚀 OpenCaselist MCP Server v9 on port {port}")
    print(f"📚 Shard: {CASELIST_SHARD}")
    print(f"🔑 Credentials: {'✅' if TABROOM_USERNAME else '❌ missing'}")
    print(f"🌐 Base URL: {SERVER_BASE_URL}")
    print(f"⚙️  Limits: {MAX_TEAMS} teams, {MAX_FILES} files, {FILE_TIMEOUT}s/file, {CONCURRENCY} concurrent")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
